#%%
## Import Library
import copy
import openai
from azure.core.credentials import AzureKeyCredential
from azure.identity import AzureDeveloperCliCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswParameters,
    PrioritizedFields,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SemanticConfiguration,
    SemanticField,
    SemanticSettings,
    SimpleField,
    VectorSearch,
    VectorSearchAlgorithmConfiguration,
)
from azure.storage.blob import BlobServiceClient
from langchain.chains import LLMChain
from langchain.llms import AzureOpenAI 
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.chat_models import AzureChatOpenAI
from langchain.memory import CosmosDBChatMessageHistory
import openai
import os
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.azuresearch import AzureSearch

import openai
import os
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.azuresearch import AzureSearch

from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import ResourceExistsError
import json

from langchain.chains import RetrievalQA
from langchain.chains.question_answering import load_qa_chain
#from langchain.retrievers import AzureCognitiveSearchRetriever
from langdetect import detect
from langchain.prompts import PromptTemplate
import re
# Create chain to answer questions
from langchain.chains import RetrievalQA
from langchain.chains.question_answering import load_qa_chain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.chains import ConversationalRetrievalChain

# Import Azure OpenAI
from langchain.llms import AzureOpenAI 
from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage

import openai
from azure.core.credentials import AzureKeyCredential
from azure.identity import AzureDeveloperCliCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswParameters,
    PrioritizedFields,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SemanticConfiguration,
    SemanticField,
    SemanticSettings,
    SimpleField,
    VectorSearch,
    VectorSearchAlgorithmConfiguration,
)

from azure.storage.blob import BlobServiceClient
from pypdf import PdfReader
from langchain.schema import Document
import openai
import os
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.azuresearch import AzureSearch
#import textwrap
import logging
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import ResourceExistsError
import json
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from io import BytesIO

index_name = "credit-proposal"
search_service = "gptdemosearch"
search_api_key = "PcAZcXbX2hJsxMYExc2SnkMFO0D94p7Zw3Qzeu5WjYAzSeDMuR5O"
storage_service = "creditproposal"
storage_api_key = "hJ2qb//J1I1KmVeDHBpwEpnwluoJzm+b6puc5h7k+dnDSFQ0oxuh1qBz+qPB/ZT7gZvGufwRbUrN+ASto6JOCw=="
connect_str = f"DefaultEndpointsProtocol=https;AccountName={storage_service};AccountKey={storage_api_key}"

connection_string = f"DefaultEndpointsProtocol=https;AccountName={storage_service};AccountKey={storage_api_key}"
container_name = "exportdocs"

doc_intell_endpoint = "https://doc-intelligence-test.cognitiveservices.azure.com/"
doc_intell_key = "9fac3bb92b3c4ef292c20df9641c7374"

# set up openai environment
os.environ["OPENAI_API_TYPE"] = "azure"
os.environ["OPENAI_API_BASE"] = "https://pwcjay.openai.azure.com/"
os.environ["OPENAI_API_VERSION"] = "2023-05-15"
os.environ["OPENAI_API_KEY"] = "f282a661571f45a0bdfdcd295ac808e7"

os.environ["AZURE_COGNITIVE_SEARCH_SERVICE_NAME"] = search_service
os.environ["AZURE_COGNITIVE_SEARCH_API_KEY"] = search_api_key
os.environ["AZURE_INDEX_NAME"] = index_name


def create_docx(client_name, json_data):
    
    storage_service = "creditproposal"
    storage_api_key = "hJ2qb//J1I1KmVeDHBpwEpnwluoJzm+b6puc5h7k+dnDSFQ0oxuh1qBz+qPB/ZT7gZvGufwRbUrN+ASto6JOCw=="
    connection_string = f"DefaultEndpointsProtocol=https;AccountName={storage_service};AccountKey={storage_api_key}"
    container_name = "exportdocs"
    
    # Create a new Word document
    document = DocxDocument()

    title_text = "Credit Proposal for " + client_name
    title_size = 20 # Font size in points

    # Create a paragraph for the title
    title_paragraph = document.add_paragraph()

    # Add the title text to the paragraph
    title_run = title_paragraph.add_run(title_text)

    # Apply formatting to the title run
    title_run.bold = True
    title_run.font.size = Pt(title_size)

    # Set the alignment of the paragraph to the center
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Convert JSON values to section headers and paragraphs in the Word document
    for item in json_data['consolidated_text']:
        section = item['section']
        context = item['output']

        # Add the section header
        document.add_heading(section, level=1)

        # Split context into lines and check each line
        for line in context.split('\n'):
            # Create a new paragraph for each line
            paragraph = document.add_paragraph()

            # Search for the pattern [RM please ... ] using regex
            matches = re.findall(r'\[RM .*?\]', line)

            if matches:
                # If there's a match, split line into parts
                parts = re.split(r'(\[RM .*?\])', line)

                for part in parts:
                    run = paragraph.add_run(part)

                    if part in matches:
                        # This part should be colored red
                        run.font.color.rgb = RGBColor(255, 0, 0)  # RGB values for red
            else:
                # Normal text
                run = paragraph.add_run(line)

    blob_name = client_name + '_Word_proposal.docx'
    
    # Save the Word document to a BytesIO object
    document_bytes = BytesIO()
    document.save(document_bytes)
    document_bytes.seek(0)  # Reset the stream position to the beginning

    # Store the Word document in Azure Blob Storage
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    container_client = blob_service_client.get_container_client(container_name)
    blob_client = container_client.get_blob_client(blob_name)
    blob_client.upload_blob(document_bytes)
    
    return 
